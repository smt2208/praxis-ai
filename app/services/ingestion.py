"""
app/services/ingestion.py

Document ingestion pipeline:
  1. Download/load file from URL or path
  2. Parse document — smart two-path strategy:
       a. Fast local extraction (PyMuPDF / python-docx) — milliseconds, zero cost
       b. LlamaParse cloud API with fast_mode=True — only when local fails
  3. Chunk text with LangChain splitter
  4. Generate hybrid embeddings (dense + sparse)
  5. Upsert into Qdrant

Called by the POST /api/v1/ingest endpoint.
"""
import os
import logging
import tempfile
import asyncio
from pathlib import Path

import httpx
from langsmith import traceable
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_qdrant import QdrantVectorStore, RetrievalMode, FastEmbedSparse

from app.config import get_settings

logger = logging.getLogger(__name__)

settings = get_settings()

MAX_DOWNLOAD_BYTES = 50 * 1024 * 1024


# --- Step 1: Download file from URL ------------------------------------

async def download_file(url: str) -> Path:
    """Download a remote file to a temp path. Returns the local path."""
    suffix = Path(url.split("?")[0]).suffix or ".pdf"
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    async with httpx.AsyncClient(timeout=60) as client:
        async with client.stream("GET", url) as response:
            response.raise_for_status()
            content_type = response.headers.get("content-type", "").split(";")[0].strip().lower()
            if content_type.startswith("text/html") or content_type.startswith("application/xhtml+xml"):
                raise ValueError(f"Unsupported content type: {content_type}")

            downloaded = 0
            async for chunk in response.aiter_bytes():
                downloaded += len(chunk)
                if downloaded > MAX_DOWNLOAD_BYTES:
                    raise ValueError("Remote file is too large to ingest safely.")
                tmp.write(chunk)
    tmp.close()
    return Path(tmp.name)


# --- Step 2a: Fast local extraction (no API call) ----------------------

def _try_fast_local_extract(file_path: Path) -> list[str] | None:
    """
    Attempt zero-latency local text extraction for supported formats.
    Returns a list of page-level strings, or None if local extraction is
    not possible (e.g. scanned PDF with no embedded text, PPTX, etc.).
    """
    ext = file_path.suffix.lower()

    if ext in (".txt", ".md"):
        text = file_path.read_text(encoding="utf-8", errors="replace")
        return [text] if text.strip() else None

    if ext == ".pdf":
        try:
            import pymupdf
            doc = pymupdf.open(str(file_path))
            pages = []
            for page in doc:
                text = page.get_text()
                if text and text.strip():
                    pages.append(text.strip())
            doc.close()
            if pages:
                return pages
        except ImportError:
            logger.debug("[ingestion] PyMuPDF not installed, skipping local PDF extraction.")
        except Exception as exc:
            logger.warning("[ingestion] PyMuPDF local extraction error: %s", exc)

    if ext == ".docx":
        try:
            from docx import Document as DocxDocument
            docx = DocxDocument(str(file_path))
            full_text = "\n".join(p.text for p in docx.paragraphs if p.text.strip())
            return [full_text] if full_text.strip() else None
        except ImportError:
            logger.debug("[ingestion] python-docx not installed, skipping DOCX extraction.")
        except Exception as exc:
            logger.warning("[ingestion] python-docx extraction error: %s", exc)

    return None


# --- Step 2b: LlamaParse cloud extraction (fallback) -------------------

def _llamaparse_extract(file_path: Path) -> list[str]:
    """Use LlamaParse with fast_mode=True to extract text from complex or scanned docs."""
    api_key = (settings.llama_cloud_api_key or "").strip()
    if not api_key or api_key.startswith("llx-...") or len(api_key) < 10:
        logger.warning("[ingestion] LLAMA_CLOUD_API_KEY is not configured or is a placeholder.")
        return []

    try:
        from llama_parse import LlamaParse
        parser = LlamaParse(
            api_key=api_key,
            result_type="markdown",
            fast_mode=True,
            verbose=False,
            num_workers=4,
        )
        documents = parser.load_data(str(file_path))
        return [doc.text.strip() for doc in documents if doc.text and doc.text.strip()]
    except Exception as exc:
        logger.error("[ingestion] LlamaParse cloud extraction failed: %s", exc)
        return []


# --- Step 2: Unified parse dispatcher ----------------------------------

@traceable(name="Parse Document", run_type="parser")
def parse_document(file_path: Path) -> list[str]:
    """
    Smart two-path parser with resilient fallback:
      1. Try fast local extraction first (PyMuPDF / python-docx / txt) - zero latency & cost.
      2. If local extraction succeeded with substantial text, return immediately.
      3. If local extraction yielded no text (scanned PDF, PPTX, etc.), fall back to LlamaParse.
      4. If LlamaParse succeeds, return cloud-parsed markdown.
      5. If LlamaParse fails/unconfigured but partial local text exists, use local text.
      6. If both produce 0 text, raise a clear descriptive ValueError.
    """
    local_pages = _try_fast_local_extract(file_path)
    if local_pages and len(local_pages) > 0:
        total_chars = sum(len(p) for p in local_pages)
        if total_chars >= 50:
            logger.info("[ingestion] Fast local extraction succeeded (%d pages, %d chars).", len(local_pages), total_chars)
            return local_pages

    logger.info("[ingestion] Local extraction yielded minimal or no text — falling back to LlamaParse cloud parser.")
    cloud_pages = _llamaparse_extract(file_path)
    if cloud_pages:
        logger.info("[ingestion] LlamaParse extraction succeeded (%d pages).", len(cloud_pages))
        return cloud_pages

    # Fallback: If LlamaParse failed or returned no text, but we had some local text
    if local_pages:
        logger.warning("[ingestion] LlamaParse unavailable or empty; falling back to %d locally extracted page(s).", len(local_pages))
        return local_pages

    raise ValueError(
        f"Could not extract readable text from '{file_path.name}'. "
        "If this is a scanned document or image-only PDF, please verify that LLAMA_CLOUD_API_KEY is configured and active."
    )


# --- Step 3: Chunk text -----------------------------------------------

@traceable(name="Chunk Texts", run_type="chain")
def chunk_texts(pages: list[str], source: str, user_id: str, conversation_id: str) -> list[Document]:
    """Split page texts into overlapping chunks with per-conversation metadata."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=100,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    all_docs = []
    for page_num, text in enumerate(pages):
        chunks = splitter.split_text(text)
        for chunk in chunks:
            all_docs.append(Document(
                page_content=chunk,
                metadata={"source": source, "page": page_num + 1, "user_id": user_id, "conversation_id": conversation_id},
            ))
    return all_docs


# --- Step 4 + 5: Embed and upsert to Qdrant --------------------------

@traceable(name="Store Vector Documents", run_type="retriever")
def store_documents(docs: list[Document], collection_name: str) -> int:
    """Generate hybrid embeddings and upsert into Qdrant. Returns count stored."""
    QdrantVectorStore.from_documents(
        documents=docs,
        embedding=OpenAIEmbeddings(
            model="text-embedding-3-small",
            api_key=settings.openai_api_key,
        ),
        sparse_embedding=FastEmbedSparse(model_name="Qdrant/bm25"),
        collection_name=collection_name,
        url=settings.qdrant_url,
        api_key=settings.qdrant_api_key,
        retrieval_mode=RetrievalMode.HYBRID,
        vector_name="dense",
        sparse_vector_name="sparse",
        force_recreate=False,
    )
    return len(docs)


# --- Main pipeline entrypoint -----------------------------------------

@traceable(name="Ingest Document Pipeline", run_type="chain")
async def ingest_document(source_url: str, user_id: str, conversation_id: str) -> int:
    """
    Full pipeline: download → parse → chunk → embed → store.
    Every chunk is tagged with user_id and conversation_id for per-conversation isolation.
    Returns the number of chunks stored.
    """
    target_collection = settings.qdrant_collection_name

    is_remote = source_url.startswith("http://") or source_url.startswith("https://")
    if is_remote:
        file_path = await download_file(source_url)
        cleanup = True
    else:
        file_path = Path(source_url)
        if not file_path.exists():
            raise FileNotFoundError(f"Local file not found: {source_url}")
        cleanup = False

    try:
        logger.info("[ingestion] Parsing: %s", file_path.name)
        pages = await asyncio.to_thread(parse_document, file_path)
        if not pages:
            raise ValueError("Parser returned no text from the document.")

        docs = await asyncio.to_thread(
            chunk_texts, pages, source_url, user_id, conversation_id
        )
        logger.info("[ingestion] Created %d chunks from %d pages.", len(docs), len(pages))

        count = await asyncio.to_thread(store_documents, docs, target_collection)
        logger.info("[ingestion] Stored %d chunks in '%s' for user '%s'.", count, target_collection, user_id)
        return count

    finally:
        if cleanup and file_path.exists():
            os.unlink(file_path)
