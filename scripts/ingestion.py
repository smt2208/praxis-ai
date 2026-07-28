"""
ingestion.py

Document ingestion pipeline:
  1. Download/load file from URL or path
  2. Parse document — smart two-path strategy:
       a. Fast local extraction (PyMuPDF / python-docx) — milliseconds, zero cost
       b. LlamaParse cloud API with fast_mode=True — only when local fails
  3. Chunk text with LangChain splitter
  4. Generate hybrid embeddings (dense + sparse)
  5. Upsert into Qdrant

Called by the POST /api/v1/ingest endpoint.
"""
import os
import tempfile
import asyncio
from pathlib import Path

import httpx
from langsmith import traceable
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_qdrant import QdrantVectorStore, RetrievalMode, FastEmbedSparse

from app.config import get_settings

settings = get_settings()

MAX_DOWNLOAD_BYTES = 50 * 1024 * 1024


# --- Step 1: Download file from URL ------------------------------------

async def download_file(url: str) -> Path:
    """Download a remote file to a temp path. Returns the local path."""
    suffix = Path(url.split("?")[0]).suffix or ".pdf"
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    async with httpx.AsyncClient(timeout=60) as client:
        async with client.stream("GET", url) as response:
            response.raise_for_status()
            content_type = response.headers.get("content-type", "").split(";")[0].strip().lower()
            if content_type.startswith("text/html") or content_type.startswith("application/xhtml+xml"):
                raise ValueError(f"Unsupported content type: {content_type}")

            downloaded = 0
            async for chunk in response.aiter_bytes():
                downloaded += len(chunk)
                if downloaded > MAX_DOWNLOAD_BYTES:
                    raise ValueError("Remote file is too large to ingest safely.")
                tmp.write(chunk)
    tmp.close()
    return Path(tmp.name)


# --- Step 2a: Fast local extraction (no API call) ----------------------

def _try_fast_local_extract(file_path: Path) -> list[str] | None:
    """
    Attempt zero-latency local text extraction for supported formats.
    Returns a list of page-level strings, or None if local extraction is
    not possible (e.g. scanned PDF with no embedded text, PPTX, etc.).

    Why: LlamaParse makes a blocking cloud API call that takes 5-30 seconds.
    For a plain-text PDF or DOCX with selectable text, local extraction
    runs in milliseconds with no network round-trip.
    """
    ext = file_path.suffix.lower()

    # Plain text / Markdown — trivially fast
    if ext in (".txt", ".md"):
        text = file_path.read_text(encoding="utf-8", errors="replace")
        return [text] if text.strip() else None

    # PDF — use PyMuPDF (C-level, no network, extracts in <200ms)
    if ext == ".pdf":
        try:
            import pymupdf  # pip install pymupdf
            doc = pymupdf.open(str(file_path))
            pages = []
            total_pages = doc.page_count
            for page in doc:
                text = page.get_text()
                if text.strip():
                    pages.append(text)
            doc.close()
            # If the PDF has selectable text on at least 60% of pages, use it.
            # Otherwise it's likely a scanned image-PDF → fall through to LlamaParse.
            if pages and (total_pages == 0 or len(pages) / total_pages >= 0.6):
                return pages
        except ImportError:
            pass  # pymupdf not installed → fall through to LlamaParse
        except Exception:
            pass  # corrupt / encrypted → fall through to LlamaParse

    # DOCX — use python-docx (pure Python, no network)
    if ext == ".docx":
        try:
            from docx import Document as DocxDocument
            docx = DocxDocument(str(file_path))
            full_text = "\n".join(p.text for p in docx.paragraphs if p.text.strip())
            return [full_text] if full_text.strip() else None
        except ImportError:
            pass
        except Exception:
            pass

    # PPTX and other complex formats → let LlamaParse handle them
    return None


# --- Step 2b: LlamaParse cloud extraction (fallback) -------------------

def _llamaparse_extract(file_path: Path) -> list[str]:
    """
    Use LlamaParse with fast_mode=True to extract text from complex docs
    (scanned PDFs, PPTX, multi-column layouts, tables, etc.).

    fast_mode=True uses the lightweight 'Fast' tier:
    - Skips heavy OCR / LLM-based layout reconstruction.
    - Dramatically faster for text-based documents.
    - Still handles tables and headings better than naive extraction.
    """
    from llama_parse import LlamaParse
    parser = LlamaParse(
        api_key=settings.llama_cloud_api_key,
        result_type="markdown",
        fast_mode=True,      # Use the Fast tier — skips heavy OCR/LLM layout pass
        verbose=False,
        num_workers=4,       # Parallelise multi-page jobs server-side
    )
    documents = parser.load_data(str(file_path))
    return [doc.text for doc in documents if doc.text.strip()]


# --- Step 2: Unified parse dispatcher ----------------------------------

@traceable(name="Parse Document", run_type="parser")
def parse_document(file_path: Path) -> list[str]:
    """
    Smart two-path parser:
      1. Try fast local extraction first (milliseconds, no API cost).
      2. Fall back to LlamaParse (seconds) only when local extraction fails
         or returns no text (scanned images, complex PPTX, etc.).
    """
    pages = _try_fast_local_extract(file_path)
    if pages:
        print(f"[ingestion] Fast local extraction succeeded ({len(pages)} pages).")
        return pages

    print("[ingestion] Local extraction insufficient — falling back to LlamaParse (fast_mode).")
    return _llamaparse_extract(file_path)


# --- Step 3: Chunk text -----------------------------------------------

@traceable(name="Chunk Texts", run_type="chain")
def chunk_texts(pages: list[str], source: str, user_id: str, conversation_id: str) -> list[Document]:
    """
    Split page texts into overlapping chunks.
    Each chunk carries source + user_id + conversation_id metadata for per-conversation isolation in Qdrant.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=100,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    all_docs = []
    for page_num, text in enumerate(pages):
        chunks = splitter.split_text(text)
        for chunk in chunks:
            all_docs.append(Document(
                page_content=chunk,
                metadata={"source": source, "page": page_num + 1, "user_id": user_id, "conversation_id": conversation_id},
            ))
    return all_docs


# --- Step 4 + 5: Embed and upsert to Qdrant --------------------------

@traceable(name="Store Vector Documents", run_type="retriever")
def store_documents(docs: list[Document], collection_name: str) -> int:
    """
    Generate hybrid embeddings and upsert into Qdrant.
    Returns the number of documents stored.
    """
    QdrantVectorStore.from_documents(
        documents=docs,
        embedding=OpenAIEmbeddings(
            model="text-embedding-3-small",
            api_key=settings.openai_api_key,
        ),
        sparse_embedding=FastEmbedSparse(model_name="Qdrant/bm25"),
        collection_name=collection_name,
        url=settings.qdrant_url,
        api_key=settings.qdrant_api_key,
        retrieval_mode=RetrievalMode.HYBRID,
        vector_name="dense",
        sparse_vector_name="sparse",
        force_recreate=False,   # Never wipe existing data
    )
    return len(docs)


# --- Main pipeline entrypoint -----------------------------------------

@traceable(name="Ingest Document Pipeline", run_type="chain")
async def ingest_document(
    source_url: str,
    user_id: str,
    conversation_id: str,
) -> int:
    """
    Full pipeline: download → parse → chunk → embed → store.
    Every chunk is tagged with user_id and conversation_id so retrieval can filter per-conversation.
    Returns the number of chunks stored.
    """
    target_collection = settings.qdrant_collection_name

    # Download file
    is_remote = source_url.startswith("http://") or source_url.startswith("https://")
    if is_remote:
        file_path = await download_file(source_url)
        cleanup = True
    else:
        file_path = Path(source_url)
        if not file_path.exists():
            raise FileNotFoundError(f"Local file not found: {source_url}")
        cleanup = False

    try:
        # Parse (blocking network/CPU call → send to thread)
        print(f"[ingestion] Parsing: {file_path.name}")
        pages = await asyncio.to_thread(parse_document, file_path)
        if not pages:
            raise ValueError("Parser returned no text from the document.")

        # Chunk (CPU bound → send to thread)
        docs = await asyncio.to_thread(
            chunk_texts, pages, source_url, user_id, conversation_id
        )
        print(f"[ingestion] Created {len(docs)} chunks from {len(pages)} pages.")

        # Store (blocking network calls to OpenAI & Qdrant → send to thread)
        count = await asyncio.to_thread(store_documents, docs, target_collection)
        print(f"[ingestion] Stored {count} chunks in '{target_collection}' for user '{user_id}'.")
        return count

    finally:
        if cleanup and file_path.exists():
            os.unlink(file_path)
